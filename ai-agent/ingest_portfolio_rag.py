"""
ingest_portfolio_rag.py - Portfolio RAG ingestion using ByteDance Dolphin

This script creates a FAISS index from portfolio documents including:
- Resume/CV
- Project READMEs
- Any markdown, text, PDF, or Word files in the data directory

Based on the reference implementation from revola-ai-agent
"""

import os
import json
import uuid
import hashlib
import logging
import argparse
import base64
from pathlib import Path
from typing import Any, Dict, List, Optional

# Initialize logger
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

# Document Processing (optional dependencies)
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. .docx files will not be processed.")

try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False
    logger.warning("beautifulsoup4 not installed. HTML files will not be processed.")

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    logger.warning("markdown not installed. .md files will be processed as plain text.")

try:
    from dolphin_parser import parse_pdf_with_dolphin, DOLPHIN_AVAILABLE
    PDF_AVAILABLE = DOLPHIN_AVAILABLE
    if not PDF_AVAILABLE:
        logger.warning("Dolphin not available. .pdf files will not be processed.")
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("dolphin_parser module not found. .pdf files will not be processed.")

# GitHub API
try:
    from github import Github
    GITHUB_AVAILABLE = True
except ImportError:
    GITHUB_AVAILABLE = False
    logger.warning("PyGithub not installed. GitHub repository ingestion will not be available.")

# LangChain and AI Models
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer

# Vector Database
import faiss
import numpy as np

# Environment
from dotenv import load_dotenv

load_dotenv()


def calculate_file_hash(filepath):
    """Calculates the SHA256 hash of a file."""
    hasher = hashlib.sha256()
    try:
        with open(filepath, "rb") as file:
            while chunk := file.read(8192):
                hasher.update(chunk)
        return hasher.hexdigest()
    except (FileNotFoundError, IOError) as e:
        logging.warning(f"Could not read file {filepath} for hashing: {e}")
        return None


def store_metadata_as_json(metadata_list, output_filepath):
    """Stores a list of metadata dictionaries to a JSON file."""
    os.makedirs(os.path.dirname(output_filepath), exist_ok=True)
    try:
        with open(output_filepath, "w", encoding="utf-8") as f:
            json.dump(metadata_list, f, indent=4, ensure_ascii=False)
        logging.info(
            f"Successfully stored metadata for {len(metadata_list)} items in {output_filepath}"
        )
        return True
    except (IOError, TypeError) as e:
        logging.error(
            f"Error writing or serializing JSON metadata to {output_filepath}: {e}"
        )
        return False


def create_and_save_faiss_index(embeddings_array: np.ndarray, output_path: str):
    """Creates a FAISS index from embeddings and saves it to a file."""
    if embeddings_array is None or embeddings_array.shape[0] == 0:
        logging.warning(
            f"No embeddings provided. Skipping FAISS index creation for {output_path}."
        )
        return False

    num_vectors, dim = embeddings_array.shape
    logging.info(
        f"Creating FAISS index for {num_vectors} vectors of dimension {dim}..."
    )

    index = faiss.IndexFlatL2(dim)
    index = faiss.IndexIDMap(index)
    ids = np.arange(num_vectors)
    index.add_with_ids(embeddings_array.astype("float32"), ids)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    faiss.write_index(index, output_path)
    logging.info(f"FAISS index saved successfully to {output_path}.")
    return True


def process_markdown_file(
    filepath: str, document_id: str, content_hash: str
) -> List[Document]:
    """Extract and render Markdown files."""
    documents = []
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            md_content = f.read()

        # Convert to HTML if markdown library available, otherwise use raw text
        if MARKDOWN_AVAILABLE:
            html = markdown.markdown(md_content, extensions=["tables", "fenced_code"])
            # Parse HTML to extract clean text
            if HTML_AVAILABLE:
                soup = BeautifulSoup(html, "html.parser")
                text = soup.get_text(separator="\n", strip=True)
            else:
                text = md_content
        else:
            text = md_content

        if text and len(text.strip()) > 40:
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "parent_document_id": document_id,
                        "parent_content_hash": content_hash,
                        "source_type": "markdown_text",
                        "source_file": filepath,
                    },
                )
            )

        logging.info(f"Extracted text from Markdown file: {filepath}")
    except Exception as e:
        logging.error(f"Error processing Markdown file {filepath}: {e}")

    return documents


def process_txt_file(
    filepath: str, document_id: str, content_hash: str
) -> List[Document]:
    """Extract text from plain text files."""
    documents = []
    try:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()

        if text and len(text.strip()) > 40:
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "parent_document_id": document_id,
                        "parent_content_hash": content_hash,
                        "source_type": "text_file",
                        "source_file": filepath,
                    },
                )
            )

        logging.info(f"Extracted text from plain text file: {filepath}")
    except Exception as e:
        logging.error(f"Error processing text file {filepath}: {e}")

    return documents


def process_docx_file(
    filepath: str, document_id: str, content_hash: str
) -> List[Document]:
    """Extract text from Word documents (.docx)."""
    if not DOCX_AVAILABLE:
        return []

    documents = []
    try:
        doc = DocxDocument(filepath)

        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Combine text
        full_text = "\n\n".join(paragraphs)

        if full_text:
            documents.append(
                Document(
                    page_content=full_text,
                    metadata={
                        "parent_document_id": document_id,
                        "parent_content_hash": content_hash,
                        "source_type": "docx_text",
                        "source_file": filepath,
                    },
                )
            )

        logging.info(f"Extracted {len(documents)} sections from Word document")
    except Exception as e:
        logging.error(f"Error processing Word document {filepath}: {e}")

    return documents


def process_pdf_file(
    filepath: str, document_id: str, content_hash: str
) -> List[Document]:
    """Extract text from PDF documents (.pdf) using ByteDance Dolphin."""
    if not PDF_AVAILABLE:
        return []

    documents = []
    try:
        # Use Dolphin to parse the PDF
        logging.info(f"Using Dolphin to parse PDF: {filepath}")
        full_text = parse_pdf_with_dolphin(filepath, model_path="./hf_model")

        if full_text and len(full_text.strip()) > 40:
            # Count pages from the output (rough estimate based on page markers)
            num_pages = full_text.count("=== Page ")

            documents.append(
                Document(
                    page_content=full_text,
                    metadata={
                        "parent_document_id": document_id,
                        "parent_content_hash": content_hash,
                        "source_type": "pdf_text_dolphin",
                        "source_file": filepath,
                        "num_pages": num_pages,
                        "parser": "dolphin-1.5",
                    },
                )
            )
            logging.info(f"Extracted text from {num_pages} pages in PDF using Dolphin: {filepath}")
        else:
            logging.warning(f"No text content found in PDF: {filepath}")

    except Exception as e:
        logging.error(f"Error processing PDF file with Dolphin {filepath}: {e}")

    return documents


def get_user_repos(
    username: str,
    github_token: Optional[str] = None,
) -> List[str]:
    """
    Fetch all public repository names for a GitHub user.

    Args:
        username: GitHub username
        github_token: Optional GitHub token for higher rate limits

    Returns:
        List of repository names in format "owner/repo"
    """
    if not GITHUB_AVAILABLE:
        logger.error("PyGithub is not installed. Cannot fetch user repositories.")
        return []

    repo_names = []

    try:
        # Initialize GitHub API
        if github_token:
            g = Github(github_token)
            logger.info(f"Using authenticated GitHub API (higher rate limits)")
        else:
            g = Github()
            logger.info(f"Using unauthenticated GitHub API (60 requests/hour limit)")

        # Get user
        logger.info(f"Fetching repositories for user: {username}")
        user = g.get_user(username)

        # Get all public repositories
        repos = user.get_repos(type='public')

        for repo in repos:
            repo_names.append(repo.full_name)
            logger.info(f"  Found repository: {repo.full_name}")

        logger.info(f"Found {len(repo_names)} public repositories for {username}")

    except Exception as e:
        logger.error(f"Error fetching repositories for user {username}: {e}")

    return repo_names


def process_github_repo(
    repo_name: str,
    github_token: Optional[str] = None,
    file_extensions: List[str] = None,
    exclude_patterns: List[str] = None,
) -> List[Document]:
    """
    Fetch and process files from a GitHub repository without cloning.

    Args:
        repo_name: Repository name in format "owner/repo"
        github_token: Optional GitHub token for higher rate limits
        file_extensions: List of file extensions to include (e.g., ['.py', '.js', '.md'])
        exclude_patterns: List of path patterns to exclude (e.g., ['test/', 'node_modules/'])

    Returns:
        List of Document objects
    """
    if not GITHUB_AVAILABLE:
        logger.error("PyGithub is not installed. Cannot process GitHub repositories.")
        return []

    if file_extensions is None:
        file_extensions = ['.py', '.js', '.jsx', '.ts', '.tsx', '.md', '.txt', '.java', '.cpp', '.c', '.go', '.rs']

    if exclude_patterns is None:
        exclude_patterns = [
            'node_modules/', 'venv/', '.venv/', '__pycache__/',
            'dist/', 'build/', '.git/', 'vendor/', 'target/'
        ]

    documents = []

    try:
        # Initialize GitHub API
        if github_token:
            g = Github(github_token)
            logger.info(f"Using authenticated GitHub API (higher rate limits)")
        else:
            g = Github()
            logger.info(f"Using unauthenticated GitHub API (60 requests/hour limit)")

        # Get repository
        logger.info(f"Fetching repository: {repo_name}")
        repo = g.get_repo(repo_name)

        # Recursive function to process repository contents
        def process_contents(contents, path=""):
            nonlocal documents

            for content_file in contents:
                # Skip excluded patterns
                if any(pattern in content_file.path for pattern in exclude_patterns):
                    logger.debug(f"Skipping excluded path: {content_file.path}")
                    continue

                if content_file.type == "dir":
                    # Recursively process directories
                    try:
                        process_contents(repo.get_contents(content_file.path), content_file.path)
                    except Exception as e:
                        logger.warning(f"Error accessing directory {content_file.path}: {e}")
                else:
                    # Process files with matching extensions
                    if any(content_file.path.endswith(ext) for ext in file_extensions):
                        try:
                            # Decode file content
                            file_content = base64.b64decode(content_file.decoded_content).decode('utf-8')

                            # Skip very large files (> 1MB)
                            if len(file_content) > 1_000_000:
                                logger.warning(f"Skipping large file (>1MB): {content_file.path}")
                                continue

                            # Skip binary files or files with minimal content
                            if len(file_content.strip()) < 10:
                                continue

                            # Create document
                            documents.append(
                                Document(
                                    page_content=file_content,
                                    metadata={
                                        "source_type": "github_file",
                                        "source_file": content_file.path,
                                        "repo_name": repo_name,
                                        "repo_url": repo.html_url,
                                        "file_url": content_file.html_url,
                                        "file_size": content_file.size,
                                        "sha": content_file.sha,
                                    },
                                )
                            )
                            logger.info(f"Processed: {content_file.path} ({content_file.size} bytes)")

                        except UnicodeDecodeError:
                            logger.warning(f"Skipping binary file: {content_file.path}")
                        except Exception as e:
                            logger.error(f"Error processing file {content_file.path}: {e}")

        # Start processing from repository root
        root_contents = repo.get_contents("")
        process_contents(root_contents)

        logger.info(f"Successfully processed {len(documents)} files from {repo_name}")

        # Get repository metadata for context
        repo_info = f"""
Repository: {repo.full_name}
Description: {repo.description or 'No description'}
Stars: {repo.stargazers_count}
Language: {repo.language or 'Multiple'}
Topics: {', '.join(repo.get_topics()) if repo.get_topics() else 'None'}
URL: {repo.html_url}
"""

        # Add repository README as first document if available
        try:
            readme = repo.get_readme()
            readme_content = base64.b64decode(readme.decoded_content).decode('utf-8')
            documents.insert(0, Document(
                page_content=f"{repo_info}\n\n{readme_content}",
                metadata={
                    "source_type": "github_readme",
                    "source_file": readme.path,
                    "repo_name": repo_name,
                    "repo_url": repo.html_url,
                },
            ))
            logger.info(f"Added README from {repo_name}")
        except Exception as e:
            logger.warning(f"Could not fetch README: {e}")

    except Exception as e:
        logger.error(f"Error accessing GitHub repository {repo_name}: {e}")

    return documents


def add_semantic_tags(chunk_text: str) -> tuple[str, dict]:
    """
    Analyzes a text chunk and prepends semantic category tags.
    Returns (enriched_text, extracted_metadata_dict)
    """
    import re

    chunk_lower = chunk_text.lower()
    categories = []
    extracted_data = {}

    # PROJECT detection
    project_patterns = [
        r"(?:project|repository|built|developed|created|implemented)",
        r"(?:github|git|npm|pip)",
        r"(?:technology|stack|framework|library)",
    ]
    if any(re.search(pattern, chunk_lower) for pattern in project_patterns):
        categories.append("PROJECT")

    # EXPERIENCE detection
    experience_patterns = [
        r"(?:experience|worked|developed|led|managed|team)",
        r"(?:role|position|job|internship|engineer|developer)",
        r"(?:company|organization|startup)",
    ]
    if any(re.search(pattern, chunk_lower) for pattern in experience_patterns):
        categories.append("EXPERIENCE")

    # SKILLS detection
    skills_patterns = [
        r"(?:python|javascript|typescript|react|node|java|c\+\+)",
        r"(?:skill|proficient|experience with|familiar with)",
        r"(?:aws|docker|kubernetes|database|api)",
    ]
    if any(re.search(pattern, chunk_lower) for pattern in skills_patterns):
        categories.append("SKILLS")

    # EDUCATION detection
    education_patterns = [
        r"(?:university|college|school|degree|bachelor|master|phd)",
        r"(?:education|studied|graduated|gpa)",
    ]
    if any(re.search(pattern, chunk_lower) for pattern in education_patterns):
        categories.append("EDUCATION")

    # Default to OTHER if no categories found
    if not categories:
        categories.append("OTHER")

    category_tags = " ".join([f"[{cat}]" for cat in categories])
    enriched_text = f"{category_tags} {chunk_text}"

    return enriched_text, extracted_data


def process_documents(
    source_directory: str,
    text_splitter: RecursiveCharacterTextSplitter,
) -> List[Dict]:
    """
    Process all documents in the source directory.

    Args:
        source_directory: Directory containing documents
        text_splitter: Text splitter for chunking

    Returns:
        List of processed chunks
    """
    all_text_chunks = []

    if not os.path.isdir(source_directory):
        logging.warning(
            f"Source directory not found: {source_directory}. Skipping text processing."
        )
        return []

    # Find all supported files
    supported_extensions = (
        ".md",
        ".markdown",
        ".txt",
        ".docx",
        ".pdf",
    )

    source_files = []
    for root, dirs, files in os.walk(source_directory):
        for f in files:
            if f.lower().endswith(supported_extensions):
                full_path = os.path.join(root, f)
                source_files.append(full_path)

    if not source_files:
        logging.warning(f"No supported files found in '{source_directory}'.")
        return []

    logging.info(
        f"\n--- Starting Document Processing for {len(source_files)} file(s) ---"
    )

    for filepath in source_files:
        relative_path = os.path.relpath(filepath, source_directory)
        document_id = relative_path.replace(os.sep, "_")
        content_hash = calculate_file_hash(filepath)
        if not content_hash:
            continue

        file_ext = filepath.lower()

        # Process different file types
        documents = []

        if file_ext.endswith((".md", ".markdown")):
            logging.info(f"\nProcessing Markdown file: {filepath}")
            documents = process_markdown_file(filepath, document_id, content_hash)
        elif file_ext.endswith(".txt"):
            logging.info(f"\nProcessing text file: {filepath}")
            documents = process_txt_file(filepath, document_id, content_hash)
        elif file_ext.endswith(".docx"):
            logging.info(f"\nProcessing Word document: {filepath}")
            documents = process_docx_file(filepath, document_id, content_hash)
        elif file_ext.endswith(".pdf"):
            logging.info(f"\nProcessing PDF document: {filepath}")
            documents = process_pdf_file(filepath, document_id, content_hash)

        # Process extracted documents
        if documents:
            split_docs = text_splitter.split_documents(documents)
            for doc in split_docs:
                doc.metadata["chunk_id"] = str(uuid.uuid4())
                enriched_text, extracted_data = add_semantic_tags(
                    doc.page_content
                )
                if extracted_data:
                    doc.metadata["extracted_data"] = extracted_data

                all_text_chunks.append(
                    {
                        "id": len(all_text_chunks),
                        "text": enriched_text,
                        "metadata": doc.metadata,
                    }
                )
            logging.info(f"Extracted {len(split_docs)} chunks from {document_id}")

    return all_text_chunks


def run_ingestion(
    source_dir: str,
    output_dir: str,
    embedding_model_name: str = "sentence-transformers/all-MiniLM-L6-v2",
    github_repos: List[str] = None,
    github_user: str = None,
    github_token: str = None,
    enable_github: bool = False,
):
    """
    Main ingestion function.

    Args:
        source_dir: Directory with source documents
        output_dir: Output directory for FAISS index
        embedding_model_name: Name of the embedding model
        github_repos: List of GitHub repositories to ingest (format: "owner/repo")
        github_user: GitHub username to fetch all public repos from
        github_token: Optional GitHub token for API authentication
        enable_github: Enable GitHub repository extraction (default: False)
    """
    logging.info(f"Starting portfolio RAG ingestion")
    logging.info(f"Source directory: {source_dir}")
    logging.info(f"Output directory: {output_dir}")
    if github_repos:
        logging.info(f"GitHub repositories: {', '.join(github_repos)}")
    if github_user:
        logging.info(f"GitHub user: {github_user}")

    # Initialize embedding model
    try:
        embeddings_model = SentenceTransformer(embedding_model_name)
        logging.info(f"Successfully initialized embedding model: {embedding_model_name}")
    except Exception as e:
        logging.error(f"Error initializing embedding model: {e}")
        return {"status": "error", "message": f"Model initialization failed: {e}"}

    # Initialize text splitter
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )

    # Process local documents
    all_text_chunks = process_documents(source_dir, text_splitter)

    # Process GitHub repositories (only if enabled)
    if enable_github:
        # Fetch all repos for a GitHub user if specified
        if github_user:
            logging.info(f"\n--- Fetching All Repositories for User: {github_user} ---")
            user_repos = get_user_repos(github_user, github_token)
            if user_repos:
                if github_repos:
                    # Combine user repos with explicitly specified repos
                    github_repos.extend(user_repos)
                else:
                    github_repos = user_repos

        # Process GitHub repositories
        if github_repos:
            logging.info("\n--- Processing GitHub Repositories ---")
            for repo_name in github_repos:
                logging.info(f"\nFetching repository: {repo_name}")
                github_docs = process_github_repo(
                    repo_name=repo_name,
                    github_token=github_token
                )

                if github_docs:
                    # Split GitHub documents into chunks
                    split_docs = text_splitter.split_documents(github_docs)
                    for doc in split_docs:
                        doc.metadata["chunk_id"] = str(uuid.uuid4())
                        enriched_text, extracted_data = add_semantic_tags(
                            doc.page_content
                        )
                        if extracted_data:
                            doc.metadata["extracted_data"] = extracted_data

                        all_text_chunks.append(
                            {
                                "id": len(all_text_chunks),
                                "text": enriched_text,
                                "metadata": doc.metadata,
                            }
                        )
                    logging.info(f"Added {len(split_docs)} chunks from {repo_name}")
    else:
        logging.info("\n--- GitHub ingestion disabled (use --enable-github to enable) ---")

    if not all_text_chunks:
        logging.warning("\n--- No data was processed. No index will be created. ---")
        return {"status": "error", "message": "No documents found to process"}

    # Generate embeddings
    logging.info(
        f"\n--- Generating Embeddings and FAISS Index for {len(all_text_chunks)} Text Chunks ---"
    )

    for i, chunk in enumerate(all_text_chunks):
        chunk["id"] = i

    text_for_embedding = [chunk["text"] for chunk in all_text_chunks]
    text_embeddings = embeddings_model.encode(
        text_for_embedding, show_progress_bar=True
    )

    # Save FAISS index and metadata
    os.makedirs(output_dir, exist_ok=True)

    index_path = os.path.join(output_dir, "portfolio_index.faiss")
    metadata_path = os.path.join(output_dir, "portfolio_metadata.json")

    index_saved = create_and_save_faiss_index(text_embeddings, index_path)
    meta_saved = store_metadata_as_json(all_text_chunks, metadata_path)

    if index_saved and meta_saved:
        logging.info("\n--- Ingestion Script Finished Successfully ---")
        return {
            "status": "success",
            "message": f"Processed {len(all_text_chunks)} chunks",
            "index_path": index_path,
            "metadata_path": metadata_path,
        }
    else:
        return {"status": "error", "message": "Failed to save index or metadata"}


def main():
    parser = argparse.ArgumentParser(
        description="Portfolio RAG ingestion script - Process local documents and GitHub repositories",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    parser.add_argument(
        "--source_dir",
        default="./data/documents",
        help="Local directory containing documents.",
    )
    parser.add_argument(
        "--output_dir",
        default="./data/faiss_index",
        help="Directory to save the FAISS index and metadata.",
    )
    parser.add_argument(
        "--embedding_model",
        default="sentence-transformers/all-MiniLM-L6-v2",
        help="Name of the sentence transformer model to use for embeddings.",
    )
    parser.add_argument(
        "--enable-github",
        action="store_true",
        help="Enable GitHub repository extraction (disabled by default)",
    )
    parser.add_argument(
        "--github_repos",
        nargs="+",
        help='GitHub repositories to ingest (format: "owner/repo"). Example: --github_repos "octocat/Hello-World" "torvalds/linux"',
    )
    parser.add_argument(
        "--github_user",
        help="GitHub username to fetch all public repositories from",
    )
    parser.add_argument(
        "--github_token",
        help="GitHub personal access token for higher API rate limits (optional but recommended)",
    )

    args = parser.parse_args()

    # Get GitHub token from args or environment variable
    github_token = args.github_token or os.getenv("GITHUB_TOKEN")

    result = run_ingestion(
        source_dir=args.source_dir,
        output_dir=args.output_dir,
        embedding_model_name=args.embedding_model,
        github_repos=args.github_repos,
        github_user=args.github_user,
        github_token=github_token,
        enable_github=args.enable_github,
    )

    print(f"\n{result}")


if __name__ == "__main__":
    main()
